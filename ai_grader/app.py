import os, subprocess, tempfile, json, re, requests, zipfile, io
from datetime import datetime
import pandas as pd
from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file
from models import db, User, Problem, TestCase, Submission, Quiz, QuizResult, LessonDocument, Question
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
import docx

app = Flask(__name__)
app.secret_key = 'ai_grader_ultimate_v3_2026'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///grader.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Khởi tạo db NGAY TẠI ĐÂY để các Class phía dưới không bị lỗi "db is not defined"
from models import db, User, Problem, TestCase, Submission, Quiz, QuizResult, LessonDocument

db.init_app(app)

with app.app_context():
    db.create_all()
# 1. CẤU HÌNH AI - GEMINI & DEEPSEEK

GEMINI_API_KEY = "AIzaSyCrJZ9ZUgBS_KqyzE6gs3G9iZNv9Ba71PU"
DEEPSEEK_API_KEY = "sk-720be83041494ed4b8083f67e303382b"

def call_gemini_api(prompt, force_json=False):
   
    model_name = "gemini-1.5-flash" 
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={GEMINI_API_KEY}"
    
    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {
            "response_mime_type": "application/json" if force_json else "text/plain",
            "temperature": 0.2
        }
    }
    try:
        response = requests.post(url, json=payload, timeout=30)
        if response.status_code == 200:
            return response.json()['candidates'][0]['content']['parts'][0]['text']
       
        return call_deepseek_api(prompt, force_json)
    except:
        return call_deepseek_api(prompt, force_json)

def call_deepseek_api(prompt, force_json=False):
    """Hàm gọi AI DeepSeek - Cứu tinh khi Gemini sập mạng"""
    if not DEEPSEEK_API_KEY or DEEPSEEK_API_KEY == "sk-720be...": 
        print(" Lỗi: Bạn chưa dán Key DeepSeek vào code!")
        return None
        
    url = "https://api.deepseek.com/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}"
    }
    
    data = {
        "model": "deepseek-coder", # Đổi sang bản coder cho chuẩn bài thi lập trình
        "messages": [
            {"role": "system", "content": "Bạn là giảng viên chấm thi lập trình C. RẤT QUAN TRỌNG: Chỉ trả về mảng JSON [...] nếu được yêu cầu, KHÔNG giải thích thêm."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3
    }
    
    try:
        # Timeout 180s 
        response = requests.post(url, headers=headers, json=data, timeout=180) 
        
        if response.status_code == 200:
            result = response.json()
            print(" DeepSeek cứu giá thành công!")
            return result['choices'][0]['message']['content']
        else:
            print(f"\n Cả DeepSeek cũng bay màu: {response.status_code}")
            print(f"Chi tiết: {response.text}")
            return None
            
    except requests.exceptions.Timeout:
        print("\n Lỗi: DeepSeek ngâm bài quá lâu (vượt 180s).")
        return None
    except Exception as e:
        print(f"\n Lỗi kết nối DeepSeek: {str(e)}")
        return None


# HÀM CẦU NỐI - GỌI THẰNG NÀY ĐỂ KÍCH HOẠT CHUỖI LIÊN HOÀN (GEMINI -> DEEPSEEK)
def generate_ai_response(prompt):
    """Bật force_json=True để ép nhả định dạng chuẩn"""
    return call_gemini_api(prompt, force_json=True)


# 3. ROUTE XÁC THỰC & CHUNG

@app.route('/')
def index(): # Hoặc tên hàm trang chủ của đại vương
    user = None
    # Nếu đã đăng nhập thì lôi thông tin ra
    if 'user_id' in session:
        user = User.query.get(session['user_id'])
        
    return render_template('home.html', user=user) # Gửi biến user ra ngoài

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        msv_or_user = request.form.get('student_id')
        
        # TÌM KIẾM THÔNG MINH: Tìm theo Mã SV trước, nếu không có thì tìm theo Tên đăng nhập
        user = User.query.filter_by(student_id=msv_or_user).first()
        if not user:
            user = User.query.filter_by(username=msv_or_user).first()
        
        if user and check_password_hash(user.password, request.form.get('password')):
            # Lấy tên hiển thị
            display = user.full_name if hasattr(user, 'full_name') and user.full_name else user.student_id
            
            # Cập nhật Session
            session.update({
                'user_id': user.id, 
                'role': user.role, 
                'msv': user.student_id, 
                'display_name': display
            })
            
            # Phân luồng theo Role 
            if user.role == 'Admin':
                return redirect('/teacher') 
            else:
                return redirect('/student')
                
        # Nếu sai tài khoản hoặc mật khẩu
        flash('Sai Mã Định Danh hoặc Mật khẩu!', 'danger')
        
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        msv = request.form.get('student_id')
        name = request.form.get('full_name')
        pw = request.form.get('password')
        hashed = generate_password_hash(pw)
        
        
        new_user = User(student_id=msv, username=msv, full_name=name, password=hashed, role='Student')
        
        try:
            db.session.add(new_user)
            db.session.commit()
            flash('Đăng ký thành công!', 'success')
            return redirect('/login')
        except Exception as e: 
            db.session.rollback()
            flash('Mã Sinh Viên hoặc Tên đăng nhập này đã tồn tại!', 'danger')
            
            
    return render_template('register.html')

@app.route('/logout')
def logout():
    session.clear() # Xóa sạch session (tên, id, role...)
    flash("Bạn đã đăng xuất thành công!", "info")
    return redirect('/') # ĐƯA VỀ TRANG CHỦ NGOÀI CÙNG

@app.route('/profile', methods=['GET', 'POST'])
def profile():
    if 'user_id' not in session: return redirect('/login')
    user = User.query.get(session['user_id'])
    if request.method == 'POST':
        user.full_name = request.form.get('full_name')
        user.phone = request.form.get('phone')
        user.school = request.form.get('school')
        db.session.commit()
        session['display_name'] = user.full_name if user.full_name else user.student_id
        flash('Đã cập nhật hồ sơ!', 'success')
    return render_template('profile.html', user=user)

# 4. ROUTE GIẢNG VIÊN (TEACHER)

@app.route('/teacher', methods=['GET', 'POST'])
@app.route('/teacher')
def teacher_dashboard():
    from models import LessonDocument
    if session.get('role') != 'Admin': 
        return redirect('/login')
        
    if request.method == 'POST':
        title = request.form.get('title')
        desc = request.form.get('description')
        m_score = float(request.form.get('max_score', 10))
        new_prob = Problem(title=title, description=desc, max_score=m_score)
        db.session.add(new_prob)
        db.session.commit()
        flash('Đã thêm bài tập mới!', 'success')
        return redirect('/teacher')
    
   
    problems = Problem.query.all()
    quizzes = Quiz.query.all() # Lấy dữ liệu từ bảng Quiz
    
    # --- PHẦN BỔ SUNG ĐỂ HIỂN THỊ TÀI LIỆU RAG ---
    # Lấy danh sách tài liệu y khoa AI đã học để hiện ra bảng
    documents = LessonDocument.query.all() 
    
    # Gửi cả 3 danh sách sang HTML: Code, Trắc nghiệm, và Tài liệu RAG
    return render_template('teacher.html', 
                           problems=problems, 
                           quizzes=quizzes, 
                           documents=documents)

@app.route('/teacher/dashboard')
def admin_scoreboard():
    if session.get('role') != 'Admin': return redirect('/login')
    
    try:
        code_subs = Submission.query.all()
        quiz_subs = QuizResult.query.all()
    except:
        return "Lỗi Database: Hãy kiểm tra tên bảng!"

    all_submissions = []
    
    # Xử lý Code Tự luận
    for c in code_subs:
        user = getattr(c, 'user', None)
        s_name = getattr(user, 'full_name', getattr(user, 'username', getattr(user, 'email', 'Sinh viên'))) if user else "Ẩn danh"
        
        # Lấy MAX SCORE từ đề bài (Nếu đề bị xóa thì mặc định là 10)
        max_p = c.problem.max_score if (hasattr(c, 'problem') and c.problem) else 10.0
        actual_score = getattr(c, 'final_score', getattr(c, 'score', 0))
        
        all_submissions.append({
            'id': getattr(c, 'id', 0), # THÊM ID Ở ĐÂY ĐỂ XEM LẠI CODE
            'type': 'Code Tự luận',
            'student_name': s_name,
            'student_email': getattr(user, 'email', 'N/A') if user else "N/A",
            'task_title': c.problem.title if (hasattr(c, 'problem') and c.problem) else 'Bài đã xóa',
            'score': actual_score,
            'max_score': max_p, # TÍNH ĐIỂM ĐỘNG THEO ĐỀ BÀI
            'time': getattr(c, 'created_at', getattr(c, 'submitted_at', None)),
            'status': getattr(c, 'status', 'Đã nộp'),
            'badge_color': 'primary', 'icon': 'bi-braces'
        })

    # Xử lý Trắc nghiệm AI
    for q in quiz_subs:
        user_q = None
        if hasattr(q, 'user') and q.user: user_q = q.user
        elif hasattr(q, 'user_id'): user_q = db.session.get(User, q.user_id)

        s_name = getattr(user_q, 'full_name', getattr(user_q, 'username', getattr(user_q, 'email', 'Sinh viên'))) if user_q else "Ẩn danh"
        
        # Lấy MAX SCORE của Trắc nghiệm (nếu không có cột này thì mặc định 10)
        max_q = getattr(q.quiz, 'max_score', 10.0) if (hasattr(q, 'quiz') and q.quiz) else 10.0

        all_submissions.append({
            'id': getattr(q, 'id', 0), # THÊM ID Ở ĐÂY ĐỂ XEM LẠI TRẮC NGHIỆM
            'type': 'Trắc nghiệm AI',
            'student_name': s_name,
            'student_email': getattr(user_q, 'email', 'N/A') if user_q else "N/A",
            'task_title': q.quiz.title if (hasattr(q, 'quiz') and q.quiz) else 'Đề đã xóa',
            'score': getattr(q, 'score', getattr(q, 'final_score', 0)),
            'max_score': max_q,
            'time': getattr(q, 'submitted_at', None),
            'status': 'Hoàn thành',
            'badge_color': 'purple', 'icon': 'bi-ui-radios'
        })

    from datetime import datetime
    all_submissions.sort(key=lambda x: x['time'] if x['time'] else datetime.min, reverse=True)
    return render_template('dashboard.html', submissions=all_submissions)


@app.route('/teacher/bulk_grade', methods=['GET', 'POST'])
def bulk_grade():
    if session.get('role') != 'Admin': return redirect('/login')
    
    def normalize(text):
        if not text: return ""
        return " ".join(str(text).replace('\r', '').strip().split()).lower()

    results = [] 
    problems_list = Problem.query.order_by(Problem.id.asc()).all()
    processed_keys = set() 
    
    if request.method == 'POST':
        zip_f = request.files.get('zip_file')
        if zip_f and zip_f.filename.endswith('.zip'):
            try:
                file_bytes = zip_f.read()
                with zipfile.ZipFile(io.BytesIO(file_bytes), 'r') as z:
                    all_files = [f for f in z.namelist() if f.lower().endswith('.c') 
                                 and not any(p.startswith('.') for p in f.split('/'))
                                 and '__MACOSX' not in f]
                    
                    for f_path in all_files:
                        parts = f_path.replace('\\', '/').split('/')
                        msv_raw = parts[-2] if len(parts) >= 2 else parts[-1].split('.')[0]
                        file_name = parts[-1].lower()

                        # 1. KHỚP BÀI THÔNG MINH
                        prob_matched = None
                        match_num = re.search(r'\d+', file_name)
                        if match_num:
                            idx = int(match_num.group())
                            prob_matched = Problem.query.get(idx)
                            if not prob_matched and 0 < idx <= len(problems_list):
                                prob_matched = problems_list[idx-1]
                        
                        if not prob_matched: continue

                        # 2. CHỐNG CHẤM LẶP
                        key = (msv_raw, prob_matched.id)
                        if key in processed_keys: continue
                        processed_keys.add(key)

                        # 3. CHẤM ĐIỂM VÀ SOẠN LOG SIÊU CHI TIẾT (GIỮ NGUYÊN YÊU CẦU)
                        test_cases = TestCase.query.filter_by(problem_id=prob_matched.id).all()
                        if not test_cases: continue

                        code = z.read(f_path).decode('utf-8', errors='ignore')
                        passed = 0
                        tc_logs = ["CHI TIẾT TEST CASE:"] # Bắt đầu log chi tiết
                        failed_list = []
                        
                        for i, tc in enumerate(test_cases, 1):
                            ok, out = run_c_code(code, tc.input_data)
                            actual = normalize(out)
                            expected = normalize(tc.expected_output)
                            
                            if ok and (expected == actual or expected in actual):
                                passed += 1
                                tc_logs.append(f"TC {i}: Đúng (Input: '{tc.input_data}' -> Output: '{actual}')")
                            else:
                                fail_str = f"TC {i}: Sai (Input: '{tc.input_data}' | Thực tế: '{actual}' | Kỳ vọng: '{expected}')"
                                tc_logs.append(fail_str)
                                failed_list.append(fail_str)

                        score = round((passed/len(test_cases)) * prob_matched.max_score, 1)
                        status = "Pass Toàn Bộ" if passed == len(test_cases) else f"Đúng {passed}/{len(test_cases)}"
                        
                        # 4. GỌI AI VỚI PROMPT CHI TIẾT
                        hint_ai = ""
                        if passed < len(test_cases):
                            all_fails = "\n".join(failed_list)
                            prompt = f"Phân tích lỗi code C sau:\n{code}\nLỗi tại:\n{all_fails}\nChỉ ra dòng sai, lý do và lời khuyên ngắn gọn dưới dạng JSON."
                            hint_ai = call_deepseek_api(prompt) or "Kiểm tra lại logic."
                        else:
                            hint_ai = "Logic xuất sắc!"

                        # GỘP LOG Y HỆT ẢNH CŨ
                        full_report = "\n".join(tc_logs) + "\n\nLỜI KHUYÊN DEEPSEEK:\n" + hint_ai

                        # 5. LƯU VÀO DATABASE
                        user = User.query.filter_by(student_id=msv_raw).first()
                        if user:
                            new_sub = Submission(
                                user_id=user.id, problem_id=prob_matched.id, 
                                code_content=code, status=status, final_score=score,
                                ai_hint=full_report # ĐÂY LÀ CHỖ LƯU ĐẦY ĐỦ LOG
                            )
                            db.session.add(new_sub)
                        
                        results.append({
                            "msv": msv_raw, "bai": prob_matched.title,
                            "score": score, "status": status,
                            "passed": f"{passed}/{len(test_cases)}", "hint": full_report 
                        })

                db.session.commit()
                session['last_grade_results'] = results 
                flash(f'Đã chấm xong {len(results)} bài với đầy đủ log chi tiết!', 'success')
            except Exception as e:
                db.session.rollback()
                flash(f'Lỗi: {str(e)}', 'danger')
                
    return render_template('bulk_grade.html', results=results, problems=problems_list)

@app.route('/teacher/export_excel')
def export_excel():
    from flask import send_file, session, redirect
    import pandas as pd
    from openpyxl.styles import Alignment
    import os

    # Kiểm tra quyền hạn và dữ liệu trong session
    if session.get('role') != 'Admin' or 'last_grade_results' not in session:
        return redirect('/teacher/bulk_grade')
    
    results = session['last_grade_results']
    if not results:
        return "Không có dữ liệu để xuất."

    df = pd.DataFrame(results)
    
    #  CẬP NHẬT NỘI DUNG ĐỂ LÀM BÀI MẪU CHO SINH VIÊN 
    def format_for_teaching(text):
        if not isinstance(text, str):
            return ""
        text = text.replace("Timeout", "Tự động ngưng vì thuật toán rơi vào lặp vô hạn")
        # Giữ nguyên toàn bộ nội dung khác (Chi tiết Test Case và Lời khuyên AI)
        return text

    if 'hint' in df.columns:
        df['hint'] = df['hint'].apply(format_for_teaching)
    # --- KẾT THÚC CẬP NHẬT ---
    
    # Tạo bản đồ ký hiệu Bài 1, Bài 2... dựa trên các đề bài có trong kết quả
    unique_probs = sorted(df['bai'].unique())
    prob_map = {title: f"Bài {i+1}" for i, title in enumerate(unique_probs)}
    df['bai_label'] = df['bai'].map(prob_map)
    
    # Pivot dữ liệu để mỗi sinh viên là một dòng
    df_score = df.pivot(index='msv', columns='bai_label', values='score').fillna(0)
    df_passed = df.pivot(index='msv', columns='bai_label', values='passed').fillna("0/0")
    df_hint = df.pivot(index='msv', columns='bai_label', values='hint').fillna("")
    
    # Khởi tạo DataFrame cuối cùng
    final_df = pd.DataFrame(index=df['msv'].unique())
    
    # Đổ dữ liệu vào theo từng cột Bài 1, Bài 2... đúng yêu cầu ban đầu
    for label in prob_map.values():
        if label in df_score.columns:
            final_df[f"{label} (Điểm)"] = df_score[label]
            final_df[f"{label} (Số Test Pass)"] = df_passed[label] 
            # Cột này chứa toàn bộ log CHI TIẾT TEST CASE và AI mà đại vương cần
            final_df[f"{label} (Chi tiết Test Case & Lời khuyên AI)"] = df_hint[label] 
    
    # Tính tổng điểm cuối cùng
    final_df['TỔNG ĐIỂM'] = df_score.sum(axis=1)
    final_df.reset_index(inplace=True)
    final_df.rename(columns={'index': 'Mã Sinh Viên'}, inplace=True)
    
    path = "Bang_Diem_Mau_Huong_Dan.xlsx"
    
    # Sử dụng pd.ExcelWriter để định dạng cột và xuống dòng
    with pd.ExcelWriter(path, engine='openpyxl') as writer:
        final_df.to_excel(writer, sheet_name='Bảng Điểm', index=False)
        
        # Sheet chú thích để sinh viên biết Bài 1, Bài 2 là đề bài nào
        mapping_info = pd.DataFrame(list(prob_map.items()), columns=['Tên gốc đề bài', 'Ký hiệu'])
        mapping_info.to_excel(writer, sheet_name='Chú thích', index=False)
        
        worksheet = writer.sheets['Bảng Điểm']
        
        # Thiết lập định dạng cột (Đặc biệt là cột Chi tiết rất dài)
        for col_idx, col_name in enumerate(final_df.columns, 1):
            from openpyxl.utils import get_column_letter
            col_letter = get_column_letter(col_idx)
            
            if "Chi tiết" in str(col_name):
                # Để cột này rộng hẳn ra (width = 80) cho sinh viên dễ đọc code phân tích
                worksheet.column_dimensions[col_letter].width = 80  
            elif "Mã Sinh Viên" in str(col_name):
                worksheet.column_dimensions[col_letter].width = 15
            else:
                worksheet.column_dimensions[col_letter].width = 15  
            
            # Áp dụng Wrap Text (xuống dòng) cho toàn bộ các dòng để hiện đầy đủ log
            for row_idx in range(1, len(final_df) + 2):
                cell = worksheet.cell(row=row_idx, column=col_idx)
                cell.alignment = Alignment(wrap_text=True, vertical='top')

    return send_file(path, as_attachment=True)

@app.route('/teacher/students')
def manage_students():
    if session.get('role') != 'Admin': return redirect('/login')
    return render_template('manage_students.html', students=User.query.filter_by(role='Student').all())

@app.route('/teacher/generate_tests/<int:id>', methods=['POST'])
def generate_tests(id):
    prob = Problem.query.get_or_404(id)
    prompt = f"Sinh 3 test case JSON cho bài: {prob.description}. Định dạng: [{{'input':'','output':''}}]. Chỉ trả về JSON."
    
    resp = call_gemini_api(prompt, force_json=True)
    if resp:
        try:
            # Bóc tách JSON cực mạnh
            import re
            match = re.search(r'\[.*\]', resp, re.DOTALL)
            if match:
                data = json.loads(match.group().replace("'", '"'))
                for t in data:
                    # Đại vương kiểm tra tên cột input/output trong bảng TestCase nhé
                    new_tc = TestCase(problem_id=id, input_data=str(t['input']), expected_output=str(t['output']))
                    db.session.add(new_tc)
                db.session.commit()
                flash('AI đã nặn xong Test Case!', 'success')
            else: flash('AI trả về sai định dạng!', 'warning')
        except Exception as e: 
            db.session.rollback()
            flash(f'Lỗi lưu Test: {str(e)}', 'danger')
    return redirect('/teacher')



# --- 1. Hàm Xóa Đề Bài ---
@app.route('/teacher/problem/delete/<int:id>')
def delete_problem(id):
    if session.get('role') != 'Admin': return redirect('/login')
    Problem.query.filter_by(id=id).delete()
    db.session.commit()
    flash('Đã xóa đề bài thành công!', 'success')
    return redirect('/teacher')

# --- 2. Hàm Thêm Sinh Viên ---
@app.route('/teacher/students/add', methods=['POST'])
def add_student():
    if session.get('role') != 'Admin': return redirect('/login')
    hashed = generate_password_hash(request.form.get('password'))
    stu = User(student_id=request.form.get('student_id'), password=hashed, role='Student', full_name=request.form.get('full_name'))
    db.session.add(stu); db.session.commit()
    flash('Đã thêm sinh viên mới!', 'success')
    return redirect('/teacher/students')

# --- 3. Hàm Xóa Sinh Viên ---
@app.route('/teacher/students/delete/<int:id>')
def delete_student(id):
    if session.get('role') != 'Admin': return redirect('/login')
    User.query.filter_by(id=id).delete()
    db.session.commit()
    flash('Đã xóa sinh viên khỏi hệ thống!', 'success')
    return redirect('/teacher/students')

# 5. ROUTE SINH VIÊN & CHẤM ĐIỂM C

def run_c_code(code, inp):
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, 'm.c')
        exe = os.path.join(tmp, 'm.exe' if os.name == 'nt' else 'm')
        with open(path, 'w', encoding='utf-8') as f: f.write(code)
        
        # ĐÃ FIX TẠI ĐÂY: Ép GCC dùng chuẩn C11 và liên kết thư viện Toán (-lm)
        cp = subprocess.run(['gcc', path, '-std=c11', '-lm', '-o', exe], capture_output=True, text=True)
        
        if cp.returncode != 0: return False, cp.stderr
        try:
            rp = subprocess.run([exe], input=inp, capture_output=True, text=True, timeout=2)
            return True, rp.stdout.strip()
        except: return False, "Timeout"

@app.route('/student')
def student_dashboard():
    if session.get('role') != 'Student': 
        return redirect('/login')
    
    # Lấy ID từ session
    user_id = session.get('user_id')
    
    # Truy vấn thông tin sinh viên để lấy tên (Ví dụ cột tên là full_name hoặc username)
    user = User.query.get(user_id) 
    
    problems = Problem.query.all()
    quizzes = Quiz.query.filter_by(is_active=True).all()
    
    # THÊM MỚI TẠI ĐÂY: Truy vấn toàn bộ tài liệu bài giảng
    documents = LessonDocument.query.all()
    
    # Gửi thêm biến user và documents sang HTML
    return render_template('student.html', 
                           problems=problems, 
                           quizzes=quizzes, 
                           user=user,
                           documents=documents)

@app.route('/student/documents')
def student_documents():
    if session.get('role') != 'Student':
        return redirect('/login')
    
    user = User.query.get(session.get('user_id'))
    # Lấy toàn bộ tài liệu bài giảng từ database
    documents = LessonDocument.query.all()
    
    return render_template('student_documents.html', user=user, documents=documents)

@app.route('/student/history')
def student_history():
    if 'user_id' not in session: return redirect('/login')
    subs = Submission.query.filter_by(user_id=session['user_id']).order_by(Submission.created_at.desc()).all()
    return render_template('student_history.html', submissions=subs)

@app.route('/student/problem/<int:id>', methods=['GET', 'POST'])
def solve_problem(id):
    prob = Problem.query.get_or_404(id)
    if request.method == 'POST':
        code = request.form.get('code', '')
        score, passed, hint = 0, 0, None
        test_cases = TestCase.query.filter_by(problem_id=id).all()
        if not test_cases:
            status = "Chưa có Test Case"
        else:
            for tc in test_cases:
                ok, out = run_c_code(code, tc.input_data)
                
                actual_out = str(out).strip()
                expected_out = str(tc.expected_output).strip()
                
                if ok and actual_out == expected_out: 
                    passed += 1
                else:
                 #  AI SOI LỖI: Chỉ soi lỗi khi có test case sai, và chỉ soi lỗi của test case đầu tiên sai để tránh spam AI quá nhiều
                    prompt = f"Code C:\n{code}\nInput: {tc.input_data}\nOutput mong đợi: {expected_out}\nOutput thực tế: {actual_out}\nLỗi ở đâu? Hãy gợi ý cách sửa."
                    
                    hint = ""
                    
                    print("\n" + "="*40)
                    print(" BẮT ĐẦU QUY TRÌNH SOI LỖI CODE")
                    print("TẦNG 1: Đang gọi máy chủ Google Gemini...")
                    
                    gemini_response = call_gemini_api(prompt)
                    
                    if gemini_response and "503" not in gemini_response and "Lỗi" not in gemini_response:
                        print(" KẾT QUẢ: Gemini trả lời THÀNH CÔNG!")
                        hint = f" [Gemini AI]: {gemini_response}"
                    else:
                        print(" THẤT BẠI: Gemini không phản hồi hoặc quá tải.")
                        print("TẦNG 2: Kích hoạt Lốp dự phòng DeepSeek...")
                        
                        deepseek_response = call_deepseek_api(prompt)
                        
                        if deepseek_response:
                            print(" KẾT QUẢ: DeepSeek gánh tạ THÀNH CÔNG!")
                            hint = f"🧬 [DeepSeek Backup]: Chuyển sang AI dự phòng.\n\n{deepseek_response}"
                        else:
                            print(" THẤT BẠI: DeepSeek cũng sập nốt.")
                            print("TẦNG 3: Khởi động chế độ bắt lỗi Offline (Local)...")
                            
                            offline_tip = "Hãy kiểm tra xem bạn có bị in thừa dấu cách, sai chữ hoa/thường không nhé."
                            if "Timeout" in actual_out:
                                offline_tip = "Code chạy quá thời gian! 99% bạn đã bị kẹt trong vòng lặp vô hạn."
                            elif "error" in actual_out.lower():
                                offline_tip = "Code bị lỗi cú pháp. Kiểm tra lại dấu chấm phẩy (;) hoặc ngoặc nhọn ({})."
                            elif actual_out.strip() == "":
                                offline_tip = "Code không in ra kết quả. Kiểm tra lại lệnh printf hoặc điều kiện if/else."
                            
                            hint = (
                                f"⚙️ [Offline Mode]: Trợ lý AI đang quá tải. Đã chuyển sang soi lỗi thủ công:\n\n"
                                f"🔹 Input: {tc.input_data}\n"
                                f"🔹 Output mong đợi: {expected_out}\n"
                                f"🔹 Output thực tế: {actual_out}\n\n"
                                f"💡 **Mẹo:** {offline_tip}"
                            )
                            print(f" KẾT QUẢ: Đã bắt bệnh Offline: {offline_tip}")
                    
                    print("="*40 + "\n")
                    # KẾT THÚC ĐOẠN THEO DÕI
                    
            score = round((passed / len(test_cases)) * prob.max_score, 2)
            status = "Pass Toàn Bộ" if passed == len(test_cases) else "Chưa đạt"
        sub = Submission(user_id=session['user_id'], problem_id=id, code_content=code, status=status, final_score=score, ai_hint=hint)
        db.session.add(sub); db.session.commit()
        return render_template('solve.html', problem=prob, score=score, status=status, ai_hint=hint, code_content=code)
    return render_template('solve.html', problem=prob)

@app.route('/teacher/generate_quiz', methods=['POST'])
def generate_quiz_ai():
    topic = request.form.get('topic', 'Lập trình C')
    num_questions = request.form.get('num_questions', '10')
    
    # --- LOG RA CMD ĐỂ THEO DÕI ---
    print("\n" + "🚀" * 10)
    print(f"[TIẾN TRÌNH] Bắt đầu gọi DeepSeek sinh đề...")
    print(f"[CHI TIẾT] Chủ đề: {topic} | Số câu: {num_questions}")
    print(f"[TRẠNG THÁI] Đang chờ AI phản hồi (Thời gian chờ tối đa 180s)...")
    
    # Prompt ép AI làm việc đúng khuôn khổ
    prompt = f"Tạo danh sách {num_questions} câu hỏi trắc nghiệm về {topic}. " \
             f"Cấu hình JSON trả về là một mảng nằm trong key 'questions'. " \
             f"Mỗi phần tử có: 'q', 'a', 'b', 'c', 'd', 'ans' (A/B/C/D)."

    try:
        raw_data = generate_ai_response(prompt)
        
        if not raw_data:
            print("❌ [LỖI] DeepSeek không trả về dữ liệu hoặc Timeout!")
            flash("DeepSeek không phản hồi, đại vương kiểm tra lại số dư tài khoản nhé!", "danger")
            return redirect('/teacher')

        print("✅ [OK] Đã nhận dữ liệu thô từ AI. Đang xử lý định dạng...")

        # Lọc sạch rác Markdown nếu có
        clean_data = raw_data.replace("```json", "").replace("```", "").strip()
        
        # Kiểm tra xem dữ liệu có phải JSON chuẩn không
        parsed_data = json.loads(clean_data)
        
        # Nếu AI trả về dạng {'questions': [...]}, mình lấy đúng cái mảng đó
        if isinstance(parsed_data, dict) and 'questions' in parsed_data:
            final_json = json.dumps(parsed_data['questions'])
            print(f"📌 [INFO] Đã tìm thấy key 'questions' và bóc tách thành công.")
        else:
            final_json = clean_data
            print(f"📌 [INFO] AI trả về mảng trực tiếp, giữ nguyên định dạng.")

        print("💾 [DATABASE] Đang lưu vào bảng Quiz...")
        new_quiz = Quiz(
            title=f"Trắc nghiệm: {topic}",
            questions=final_json,
            is_active=True
        )
        db.session.add(new_quiz)
        db.session.commit()
        
        print(f"🎉 [THÀNH CÔNG] Đã tạo xong bài: {topic}")
        print("🚀" * 10 + "\n")
        
        flash(f"Đã dùng DeepSeek nặn thành công {num_questions} câu về {topic}!", "success")
    except Exception as e:
        # Log lỗi chi tiết ra CMD để đại vương sửa nếu có biến
        print(f"💥 [LỖI CHI TIẾT]: {str(e)}")
        flash(f"Lỗi xử lý dữ liệu AI: {str(e)}", "danger")

    return redirect('/teacher')

@app.route('/student/submit_quiz/<int:quiz_id>', methods=['POST'])
def submit_quiz(quiz_id):
    quiz = Quiz.query.get(quiz_id)
    questions_list = json.loads(quiz.questions)
    
    correct_count = 0
    total = len(questions_list)
    
    for i, q in enumerate(questions_list):
        student_ans = request.form.get(f'question_{i}') # Lấy đáp án A, B, C hoặc D sinh viên chọn
        if student_ans == q['ans']:
            correct_count += 1
            
    # Công thức của thầy: (Số câu đúng / Tổng) * 10
    final_score = round((correct_count / total) * 10, 2)
    
    # Lưu lịch sử
    result = QuizResult(user_id=session['user_id'], quiz_id=quiz_id, score=final_score)
    db.session.add(result)
    db.session.commit()
    
    return f"Bạn đã hoàn thành! Điểm của bạn là: {final_score}/10"

@app.route('/teacher/toggle_quiz/<int:quiz_id>')
def toggle_quiz(quiz_id):
    quiz = Quiz.query.get(quiz_id)
    quiz.is_active = not quiz.is_active
    db.session.commit()
    return redirect('/teacher/manage_quizzes')

@app.route('/teacher/quiz/delete/<int:quiz_id>')
def delete_quiz(quiz_id):
    # Tìm bài trắc nghiệm cần xóa
    quiz = Quiz.query.get(quiz_id)
    if quiz:
        # Quan trọng: Xóa sạch lịch sử điểm của sinh viên liên quan đến bài này trước (tránh lỗi Foreign Key)
        QuizResult.query.filter_by(quiz_id=quiz_id).delete()
        
        # Sau đó mới xóa bài trắc nghiệm
        db.session.delete(quiz)
        db.session.commit()
        
    # Xóa xong thì load lại trang giáo viên
    return redirect('/teacher') # Trỏ về đúng route trang quản lý của đại vương

@app.route('/teacher/quiz/view/<int:quiz_id>')
def view_quiz(quiz_id):
    if session.get('role') != 'Admin': return redirect('/login')
    
    # Tìm bài trắc nghiệm theo ID
    quiz = Quiz.query.get_or_404(quiz_id)
    
    # Chuyển chuỗi JSON trong DB thành danh sách Python để hiển thị
    import json
    try:
        questions = json.loads(quiz.questions)
    except:
        questions = [] # Phòng trường hợp dữ liệu lỗi
        
    return render_template('view_quiz.html', quiz=quiz, questions=questions)

# --- 1. Hàm hiển thị đề thi (Để hết lỗi 404) ---
@app.route('/student/quiz/take/<int:quiz_id>')
def take_quiz_view(quiz_id): # Đổi tên thành take_quiz_view cho chắc
    if session.get('role') != 'Student': return redirect('/login')
    quiz = Quiz.query.get_or_404(quiz_id)
    import json
    questions = json.loads(quiz.questions)
    return render_template('take_quiz.html', quiz=quiz, questions=questions)

@app.route('/student/quiz/submit_action/<int:quiz_id>', methods=['POST'])
def submit_quiz_action(quiz_id):
    if session.get('role') != 'Student': return redirect('/login')
    quiz = Quiz.query.get_or_404(quiz_id)
    questions = json.loads(quiz.questions)
    
    score = 0
    total = len(questions)
    user_id = session.get('user_id')
    
    # 1. Thu thập đáp án sinh viên đã chọn
    user_choices = {}
    for i, q in enumerate(questions):
        choice = request.form.get(f'q_{i}')
        user_choices[str(i)] = choice
        if choice == q['ans']:
            score += (10 / total)

    # 2. Lưu vào lịch sử (QuizResult)
    new_result = QuizResult(
        user_id=user_id, 
        quiz_id=quiz.id, 
        score=round(score, 2),
        user_answers=json.dumps(user_choices) # Lưu lại để sau này xem lại bài
    )
    db.session.add(new_result)
    db.session.commit()

    # 3. Chuyển hướng sang trang "Xem lại bài vừa làm"
    flash(f"Nộp bài thành công! Điểm của bạn: {round(score, 2)}", "success")
    return redirect(f'/student/quiz/result/{new_result.id}')

# ROUTE MỚI: Xem lại kết quả chi tiết
@app.route('/student/quiz/result/<int:result_id>')
def view_quiz_result(result_id):
    if session.get('role') != 'Student': return redirect('/login')
    
    result = QuizResult.query.get_or_404(result_id)
    quiz = Quiz.query.get(result.quiz_id)
    
    questions = json.loads(quiz.questions)
    user_answers = json.loads(result.user_answers)
    
    return render_template('quiz_result.html', quiz=quiz, questions=questions, user_answers=user_answers, result=result)

@app.route('/teacher/quiz/result/<int:result_id>')
def view_quiz_result_teacher(result_id):
    if session.get('role') != 'Admin':
        return redirect('/login')
    
    result = QuizResult.query.get_or_404(result_id)
    quiz = Quiz.query.get_or_404(result.quiz_id)
    
    questions = json.loads(quiz.questions)
    user_answers = json.loads(result.user_answers)
    
    return render_template(
        'quiz_result.html',
        quiz=quiz,
        questions=questions,
        user_answers=user_answers,
        result=result,
        teacher_view=True
    )

@app.route('/student/quiz-history')
def student_quiz_history():
    if session.get('role') != 'Student': return redirect('/login')
    user_id = session.get('user_id')
    
    quiz_history = QuizResult.query.filter_by(user_id=user_id).all()
    try: code_history = Submission.query.filter_by(user_id=user_id).all()
    except: code_history = []
        
    all_history = []
    for q in quiz_history:
        max_q = getattr(q.quiz, 'max_score', 10.0) if (hasattr(q, 'quiz') and q.quiz) else 10.0
        all_history.append({
            'type': 'Trắc nghiệm AI',
            'title': q.quiz.title if (hasattr(q, 'quiz') and q.quiz) else 'Đề thi đã xóa',
            'score': getattr(q, 'score', 0),
            'max_score': max_q,
            'time': getattr(q, 'submitted_at', None),
            'badge_color': 'purple', 'icon': 'bi-ui-radios',
            'id': getattr(q, 'id', 0)
        })
        
    for c in code_history:
        max_p = c.problem.max_score if (hasattr(c, 'problem') and c.problem) else 10.0
        actual_score = getattr(c, 'final_score', getattr(c, 'score', 0))
        all_history.append({
            'type': 'Code Tự luận',
            'title': c.problem.title if (hasattr(c, 'problem') and c.problem) else 'Bài đã xóa',
            'score': actual_score,
            'max_score': max_p,
            'time': getattr(c, 'created_at', getattr(c, 'submitted_at', None)),
            'badge_color': 'primary', 'icon': 'bi-braces',
            'id': getattr(c, 'id', 0)
        })
        
    from datetime import datetime
    all_history.sort(key=lambda x: x['time'] if x['time'] else datetime.min, reverse=True)
    return render_template('student_history.html', history=all_history)



# ROUTE: SINH VIÊN XEM LẠI CODE CỦA MÌNH

@app.route('/student/submission/<int:id>')
def view_submission_student(id):
    if session.get('role') != 'Student': 
        return redirect('/login')
        
    sub = Submission.query.get_or_404(id)
    
    # Bảo mật: Không cho sinh viên xem lén bài của nhau
    if sub.user_id != session.get('user_id'):
        flash("Bạn không có quyền xem bài nộp của người khác!", "danger")
        return redirect('/student/quiz-history')
        
    
    return render_template('view_code.html', sub=sub)



# ROUTE: GIẢNG VIÊN XEM CODE SINH VIÊN

@app.route('/teacher/submission/<int:id>')
def view_submission_teacher(id):
    if session.get('role') != 'Admin': 
        return redirect('/login')
        
    sub = Submission.query.get_or_404(id)
    
    
    return render_template('view_code.html', sub=sub)


def extract_text_from_pdf(file_path):
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    except Exception as e:
        print(f"Lỗi đọc PDF: {e}")
    return text

def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Lỗi đọc DOCX: {e}")
    return text

def generate_quiz_from_document(extracted_text, num_questions=5):
    # GIỚI HẠN KÝ TỰ: Nếu tài liệu quá dài, chỉ lấy khoảng 15000 ký tự đầu để tránh AI bị "ngợp"
    safe_text = extracted_text[:15000] 
    
    prompt = f"""
    Đóng vai trò là một giảng viên y khoa. Hãy đọc nội dung TÀI LIỆU CUNG CẤP dưới đây và tạo ra {num_questions} câu hỏi trắc nghiệm.
    
    YÊU CẦU BẮT BUỘC (SỐNG CÒN): 
    1. Chỉ sử dụng kiến thức nằm TRONG tài liệu này. TUYỆT ĐỐI KHÔNG lấy thông tin từ bên ngoài.
    2. Trả về đúng định dạng JSON dạng mảng (Array), không kèm văn bản nào khác.
    Cấu trúc JSON: [{{"q": "Câu hỏi", "a": "Đáp án A", "b": "Đáp án B", "c": "Đáp án C", "d": "Đáp án D", "ans": "A/B/C/D", "explain": "Giải thích ngắn gọn tại sao đúng dựa trên tài liệu"}}]
    
    TÀI LIỆU CUNG CẤP:
    {safe_text}
    """
    


UPLOAD_FOLDER = 'static/uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx'}




def extract_text_from_pdf(file_path):
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    except Exception as e:
        print(f"Lỗi đọc PDF: {e}")
    return text

def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Lỗi đọc DOCX: {e}")
    return text

# Thiết lập thư mục lưu tài liệu bài giảng
DOC_UPLOAD_FOLDER = 'static/documents'
app.config['DOC_UPLOAD_FOLDER'] = DOC_UPLOAD_FOLDER
if not os.path.exists(DOC_UPLOAD_FOLDER):
    os.makedirs(DOC_UPLOAD_FOLDER)

ALLOWED_DOC_EXTENSIONS = {'pdf', 'docx'}

def allowed_doc_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_DOC_EXTENSIONS

# API Xử lý khi Giảng viên bấm nút Tải tài liệu lên
# API Xử lý khi Giảng viên bấm nút Tải tài liệu lên
@app.route('/upload_document', methods=['POST'])
def upload_document():
    from models import LessonDocument
    # Chiêu ép tạo bảng đã thành công, cứ giữ lại cho chắc cú!
    db.create_all() 

    title = request.form.get('title')
    file = request.files.get('file')
    
    if file and allowed_doc_file(file.filename):
        filename = secure_filename(file.filename)
        
        # Thêm thời gian vào tên file để không bị trùng lặp
        time_str = datetime.now().strftime("%Y%m%d_%H%M%S_")
        safe_filename = time_str + filename
        file_path = os.path.join(app.config['DOC_UPLOAD_FOLDER'], safe_filename)
        
        # Lưu file vật lý
        file.save(file_path)
        
        # Bóc tách chữ
        extracted_text = ""
        if filename.endswith('.pdf'):
            extracted_text = extract_text_from_pdf(file_path)
        elif filename.endswith('.docx'):
            extracted_text = extract_text_from_docx(file_path)
            
        # Lưu vào Database
        new_doc = LessonDocument(
            title=title, 
            file_name=filename, 
            file_path=file_path, 
            extracted_text=extracted_text
        )
        db.session.add(new_doc)
        db.session.commit()
        
        
        flash("Tải tài liệu thành công! AI đã học xong kiến thức của bài này.", "success")
        return redirect(url_for('teacher_dashboard'))
        
    # NẾU LỖI THÌ CŨNG PHẢI RETURN
    flash("File không hợp lệ. Hệ thống chỉ chấp nhận định dạng .pdf hoặc .docx", "danger")
    return redirect(url_for('teacher_dashboard'))

# API Yêu cầu AI sinh đề từ một tài liệu cụ thể
@app.route('/generate_quiz_from_doc/<int:doc_id>', methods=['POST'])
def generate_quiz_from_doc(doc_id):
    # 1. Lấy tham số và dữ liệu gốc
    num_questions = int(request.form.get('num_questions', 5))
    document = LessonDocument.query.get_or_404(doc_id)
    
    if not document.extracted_text or len(document.extracted_text.strip()) < 50:
        flash("Tài liệu quá ngắn hoặc chưa có dữ liệu để sinh đề!", "warning")
        return redirect(url_for('teacher_dashboard'))
        
    # 2. Chuẩn bị "nguyên liệu" cho AI DeepSeek
    safe_text = document.extracted_text[:12000] 
    
    prompt = f"""
    Bạn là một chuyên gia khảo thí y khoa cao cấp.
    Nhiệm vụ: Dựa trên tài liệu dưới đây, hãy soạn {num_questions} câu hỏi trắc nghiệm.
    
    QUY TẮC NGHIÊM NGẶT:
    1. Câu hỏi và đáp án PHẢI dựa hoàn toàn vào thông tin trong tài liệu.
    2. Độ khó: Phân loại từ dễ đến khó.
    3. Định dạng: Chỉ trả về duy nhất mã JSON theo cấu trúc mảng, không giải thích gì thêm.
    
    Cấu trúc mẫu:
    [
      {{"q": "Câu hỏi 1 là gì?", "a": "Đán án A", "b": "Đáp án B", "c": "Đáp án C", "d": "Đáp án D", "ans": "A", "explain": "Giải thích ngắn gọn"}}
    ]

    TÀI LIỆU:
    {safe_text}
    """

    try:
        # 3. GỌI API DEEPSEEK 
        raw_ai_data = call_deepseek_api(prompt, force_json=True)
        
        if not raw_ai_data:
             raise Exception("DeepSeek không trả về kết quả hoặc bị timeout.")

        # Xóa các ký tự thừa  để parse JSON cho an toàn
        raw_ai_data = raw_ai_data.strip()
        if raw_ai_data.startswith("```json"):
            raw_ai_data = raw_ai_data[7:-3]
            
        quiz_list = json.loads(raw_ai_data) 
        
        # 4. XỬ LÝ LƯU VÀO DATABASE
        # Nhét toàn bộ chuỗi JSON vào cột questions theo đúng thiết kế
        new_quiz = Quiz(
            title=f"Quizz: {document.title}", 
            questions=json.dumps(quiz_list, ensure_ascii=False)
        )
        
        db.session.add(new_quiz)
        db.session.commit()
        
        flash(f"Thành công! DeepSeek đã tạo bộ đề '{new_quiz.title}' với {len(quiz_list)} câu hỏi.", "success")
        
    except Exception as e:
        db.session.rollback() 
        flash(f"Lỗi khi DeepSeek xử lý dữ liệu: {str(e)}", "danger")
        print(f"DEBUG ERROR: {e}")

    return redirect(url_for('teacher_dashboard'))

# 4. ÉP TẠO BẢNG VÀ KHỞI TẠO TÀI KHOẢN ADMIN

with app.app_context():
    db.create_all()
    
    # Kiểm tra xem tài khoản admin đã tồn tại chưa
    admin_user = User.query.filter_by(username='admin').first()
    
    if not admin_user:
        # Nếu chưa có thì tự động tạo mới
        from werkzeug.security import generate_password_hash
        new_admin = User(
            student_id='ADMIN_01',
            username='admin', 
            password=generate_password_hash('admin123'), 
            role='Admin', 
            full_name='Giảng Viên Quản Trị'
        )
        db.session.add(new_admin)
        db.session.commit()
        print("🚀 Đã tự động tạo tài khoản: admin | Mật khẩu: admin123")

@app.route('/view_document/<int:doc_id>')
def view_document(doc_id):
    if 'user_id' not in session:
        return redirect('/login')
        
    doc = LessonDocument.query.get_or_404(doc_id)
    
    # send_file với as_attachment=False sẽ ép trình duyệt mở file (như PDF) ngay trên tab mới
    # thay vì tự động tải xuống.
    return send_file(doc.file_path, as_attachment=False)

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        # Kiểm tra xem tài khoản admin đã tồn tại chưa
        admin_user = User.query.filter_by(username='admin').first()
        
        if not admin_user:
            # Nếu chưa có thì tự động tạo mới
            new_admin = User(
                student_id='ADMIN_01',
                username='admin', 
                password=generate_password_hash('admin123'), 
                role='Admin', 
                full_name='Giảng Viên Quản Trị'
            )
            db.session.add(new_admin)
            db.session.commit()
            print(" Đã tự động tạo tài khoản: admin | Mật khẩu: admin123")
            
    app.run(debug=True)